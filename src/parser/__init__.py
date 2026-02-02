"""文档解析器基类"""

import logging
import os
import platform
import re
import sys
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, List, Type

logger = logging.getLogger(__name__)

# ============ 平台检测 ============
IS_WINDOWS = platform.system() == "Windows"


class BaseParser(ABC):
    """解析器基类"""

    @abstractmethod
    def parse(self, file_path: str) -> List[str]:
        """解析文件，返回文本段落列表"""
        pass

    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """检查是否支持该文件格式"""
        pass


class DocxParser(BaseParser):
    """Word文档解析器"""

    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in [".docx"]

    def parse(self, file_path: str) -> List[str]:
        from docx import Document

        file_size = os.path.getsize(file_path)
        print(
            f"   📄 解析Word文档: {Path(file_path).name} ({file_size / 1024:.1f} KB)",
            file=sys.stderr,
            flush=True,
        )

        doc = Document(file_path)
        paragraphs = []

        # 解析段落
        para_count = len(doc.paragraphs)
        print(f"   📝 处理 {para_count} 个段落...", file=sys.stderr, flush=True)

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格内容
        table_count = len(doc.tables)
        if table_count > 0:
            print(f"   📊 发现 {table_count} 个表格...", file=sys.stderr, flush=True)
            for table_idx, table in enumerate(doc.tables):
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)

        print(
            f"   ✅ Word解析完成: {len(paragraphs)} 个有效段落",
            file=sys.stderr,
            flush=True,
        )
        return paragraphs


class PdfParser(BaseParser):
    """PDF文档解析器"""

    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".pdf"

    def parse(self, file_path: str) -> List[str]:
        import fitz  # PyMuDF

        file_size = os.path.getsize(file_path)
        print(
            f"   📄 解析PDF文档: {Path(file_path).name} ({file_size / 1024:.1f} KB)",
            file=sys.stderr,
            flush=True,
        )

        doc = fitz.open(file_path)
        page_count = len(doc)
        print(f"   📃 PDF共 {page_count} 页...", file=sys.stderr, flush=True)

        paragraphs = []

        for page_num, page in enumerate(doc):
            if (page_num + 1) % 5 == 0 or page_num == 0:
                print(
                    f"   📄 处理第 {page_num + 1}/{page_count} 页...",
                    file=sys.stderr,
                    flush=True,
                )

            text = page.get_text()
            # 按段落分割
            for para in text.split("\n\n"):
                para = para.strip()
                if para:
                    paragraphs.append(para)

        print(
            f"   ✅ PDF解析完成: {len(paragraphs)} 个段落", file=sys.stderr, flush=True
        )
        return paragraphs


class MarkdownParser(BaseParser):
    """Markdown文档解析器"""

    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".md"

    def parse(self, file_path: str) -> List[str]:
        file_size = os.path.getsize(file_path)
        print(
            f"   📄 解析Markdown: {Path(file_path).name} ({file_size / 1024:.1f} KB)",
            file=sys.stderr,
            flush=True,
        )

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        original_length = len(content)
        print(f"   📝 原始内容 {original_length} 字符...", file=sys.stderr, flush=True)

        # 清理 Markdown 格式
        content = self._clean_markdown(content)
        cleaned_length = len(content)
        print(
            f"   🧹 清理后 {cleaned_length} 字符 (移除 {original_length - cleaned_length} 字符格式)",
            file=sys.stderr,
            flush=True,
        )

        # 分割为段落
        paragraphs = self._split_paragraphs(content)

        print(
            f"   ✅ Markdown解析完成: {len(paragraphs)} 个有效段落",
            file=sys.stderr,
            flush=True,
        )
        return paragraphs

    def _clean_markdown(self, content: str) -> str:
        """清理 Markdown 格式元素"""
        # 移除 YAML front matter
        content = re.sub(r"^---\n[\s\S]*?---\n", "", content)

        # 移除代码块
        content = re.sub(r"```[\s\S]*?```", "", content)

        # 移除行内代码
        content = re.sub(r"`[^`]+`", "", content)

        # 移除图片链接
        content = re.sub(r"!\[.*?\]\(.*?\)", "", content)

        # 移除链接，保留文本
        content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)

        return content

    def _split_paragraphs(self, content: str) -> List[str]:
        """分割并清理段落"""
        paragraphs = []
        for para in content.split("\n\n"):
            para = para.strip()

            # 移除标题符号
            para = re.sub(r"^#+\s+", "", para)

            # 移除列表符号
            para = re.sub(r"^[\s]*[-*+]\s+", "", para)
            para = re.sub(r"^[\s]*\d+\.\s+", "", para)

            if para and len(para) > 5:  # 过滤太短的段落
                paragraphs.append(para)

        return paragraphs


class ParserManager:
    """解析器管理器"""

    def __init__(self):
        self.parsers: List[BaseParser] = [
            DocxParser(),
            PdfParser(),
            MarkdownParser(),
        ]

    def parse_file(self, file_path: str) -> List[str]:
        """解析单个文件"""
        for parser in self.parsers:
            if parser.supports(file_path):
                print(
                    f"   🔍 使用 {parser.__class__.__name__} 解析 {Path(file_path).name}",
                    file=sys.stderr,
                    flush=True,
                )
                return parser.parse(file_path)
        raise ValueError(f"不支持的文件格式: {file_path}")

    def parse_directory(
        self, dir_path: str, recursive: bool = True
    ) -> Dict[str, List[str]]:
        """解析整个目录（跨平台）"""
        from pathlib import Path

        result = {}
        path = Path(dir_path)

        # 确保目录存在
        if not path.exists():
            raise ValueError(f"目录不存在: {dir_path}")

        if not path.is_dir():
            raise ValueError(f"路径不是目录: {dir_path}")

        print(f"   📂 扫描目录: {dir_path}", file=sys.stderr, flush=True)

        if recursive:
            files = list(path.rglob("*"))
        else:
            files = list(path.glob("*"))

        # 过滤只保留文件
        files = [f for f in files if f.is_file()]

        # 忽略 .gitkeep 和隐藏文件
        files = [
            f for f in files if not f.name.startswith(".") and f.name != ".gitkeep"
        ]

        print(f"   📁 找到 {len(files)} 个文件", file=sys.stderr, flush=True)

        for file_idx, file_path in enumerate(files):
            if file_idx > 0:
                print(f"   ---", file=sys.stderr, flush=True)

            try:
                content = self.parse_file(str(file_path))
                if content:
                    result[str(file_path)] = content
                    print(
                        f"   ✅ 文件处理完成: {len(content)} 个段落",
                        file=sys.stderr,
                        flush=True,
                    )
                else:
                    print(
                        f"   ⚠️ 文件无有效内容: {file_path.name}",
                        file=sys.stderr,
                        flush=True,
                    )
            except Exception as e:
                print(
                    f"   ❌ 解析失败 {file_path.name}: {e}", file=sys.stderr, flush=True
                )
                logger.warning(f"解析失败 {file_path}: {e}")

        print(
            f"   🎉 目录解析完成! 共 {len(result)} 个有效文档",
            file=sys.stderr,
            flush=True,
        )
        return result

    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        extensions = []
        for parser in self.parsers:
            if hasattr(parser, "extensions"):
                extensions.extend(parser.extensions)
        return list(set(extensions))


"""
文档解析模块

本模块提供多格式文档解析功能，支持 Word、PDF 和 Markdown 文档的解析。

## 支持的格式

| 解析器 | 文件扩展名 | 依赖库 |
|--------|-----------|--------|
| DocxParser | .docx | python-docx |
| PdfParser | .pdf | PyMuDF (pymupdf) |
| MarkdownParser | .md | 标准库 |

## 使用示例

```python
from src.parser import ParserManager

# 创建解析器
manager = ParserManager()

# 解析单个文件
paragraphs = manager.parse_file("document.docx")

# 解析整个目录
documents = manager.parse_directory("./documents", recursive=True)

# 获取支持的文件格式
extensions = manager.get_supported_extensions()
print(extensions)  # ['.docx', '.pdf', '.md']
```

## 架构设计

```
src/parser/
├── BaseParser        # 抽象基类，定义解析器接口
├── DocxParser        # Word 文档解析器
├── PdfParser         # PDF 文档解析器
├── MarkdownParser    # Markdown 解析器
└── ParserManager     # 解析器管理器
```

## 扩展开发

如需支持新的文档格式，可以：

1. 创建新的解析器类继承 BaseParser
2. 实现 parse() 和 supports() 方法
3. 在 ParserManager 中注册新解析器

```python
from src.parser import BaseParser, ParserManager

class MyParser(BaseParser):
    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.myext'
    
    def parse(self, file_path: str) -> List[str]:
        # 解析逻辑
        return paragraphs

# 注册解析器
manager = ParserManager()
manager.parsers.append(MyParser())
```

## 注意事项

- 解析器会自动检测文件类型
- 解析失败时会该文件
-记录日志并跳过 支持递归/非递归目录解析
- 支持跨平台路径处理
"""

__all__ = [
    "BaseParser",
    "DocxParser",
    "PdfParser",
    "MarkdownParser",
    "ParserManager",
]
